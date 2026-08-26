#!/usr/bin/env python3
"""Build the DailyFlow 2.2 DeepSeek Harness implementation plan DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "DAILYFLOW_2_2_DEEPSEEK_HARNESS_IMPLEMENTATION_PLAN.md"
OUTPUT = ROOT / "docs" / "DailyFlow_2.2_DeepSeek_Harness_完整开发实施计划.docx"

INK = "162338"
BLUE = "2455C3"
BLUE_DARK = "173B7A"
CYAN = "12A3A8"
MUTED = "667085"
LIGHT = "EEF3FF"
LIGHTER = "F7F9FC"
BORDER = "D7DEEA"
WHITE = "FFFFFF"
CODE_BG = "F3F5F8"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def new_numbering_instance(doc: Document, style_name: str = "List Number") -> int:
    """Create a real Word numbering instance that restarts at 1."""
    style_ppr = doc.styles[style_name]._element.pPr
    base_num_id = int(style_ppr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_run_font(run, latin="Arial Unicode MS", east_asia="Arial Unicode MS", size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str, base_size=10.5, color=INK, force_bold=False) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|<https?://[^>]+>)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=color, bold=force_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Arial Unicode MS", east_asia="Arial Unicode MS", size=max(8.5, base_size - 1), color=BLUE_DARK)
            set_cell_like_run_shading(run, "EAF0FF")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=BLUE, italic=False)
            run.font.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=color, bold=force_bold)


def set_cell_like_run_shading(run, fill: str) -> None:
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (28, BLUE_DARK, 0, 8),
        "Subtitle": (13, MUTED, 0, 10),
        "Heading 1": (16, BLUE_DARK, 18, 9),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, CYAN, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("DAILYFLOW · PRODUCT & ENGINEERING")
    set_run_font(r, size=10, color=CYAN, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(65)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("DailyFlow 2.2")
    set_run_font(r, size=34, color=BLUE_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("DeepSeek Harness 驱动的\nAI Event Operator")
    set_run_font(r, size=25, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("完整开发实施计划 · 可直接交给实现型 AI 执行")
    set_run_font(r, size=13, color=MUTED)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    data = [
        ("主 Runtime", "DeepSeek Harness（DSH）"),
        ("产品控制面", "Event · Mindmap · Evidence · Commitment · Proposal · Today"),
        ("可选子代理", "Codex via official DSH subagent plugin"),
        ("目标版本", "DailyFlow 2.2"),
        ("计划日期", "2026 年 8 月 22 日"),
    ]
    header = table.rows[0]
    for cell, text in zip(header.cells, ("字段", "内容")):
        set_cell_shading(cell, BLUE_DARK)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(header)
    for row, (label, value) in zip(table.rows[1:], data):
        set_cell_shading(row.cells[0], LIGHT)
        for cell, text, bold, color in ((row.cells[0], label, True, BLUE_DARK), (row.cells[1], value, False, INK)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=color, bold=bold)
    set_table_geometry(table, [2100, 7260])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("核心原则")
    set_run_font(r, size=10, color=CYAN, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Harness 负责运行，DailyFlow 负责产品灵魂；AI 只提议，用户确认后才写入。")
    set_run_font(r, size=14, color=INK, bold=True)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    r = hp.add_run("DAILYFLOW 2.2  /  DEEPSEEK HARNESS IMPLEMENTATION PLAN")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("DailyFlow 2.2  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(fp)


def table_widths(n: int) -> list[int]:
    if n == 2:
        return [2500, 6860]
    if n == 3:
        return [1800, 3780, 3780]
    if n == 4:
        return [1500, 2200, 2830, 2830]
    return [9360 // n] * n


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    n = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        prevent_row_split(table.rows[ri])
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            if ri == 0:
                set_cell_shading(cell, BLUE_DARK)
            elif ri % 2 == 0:
                set_cell_shading(cell, LIGHTER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, text.strip(), base_size=8.7 if n >= 4 else 9.2, color=WHITE if ri == 0 else INK, force_bold=ri == 0)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, table_widths(n))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines).rstrip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, latin="Arial Unicode MS", east_asia="Arial Unicode MS", size=8.2, color=INK)


def add_toc(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph("执行目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    lead = doc.add_paragraph()
    add_inline(lead, "本目录保留一级实施章节；具体任务以 DFH 编号领取。", color=MUTED)
    toc_num_id = new_numbering_instance(doc)
    for heading in headings:
        p = doc.add_paragraph(style="List Number")
        apply_numbering(p, toc_num_id)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        clean_heading = re.sub(r"^\d+\.\s*", "", heading)
        add_inline(p, clean_heading, base_size=9.0, color=INK)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)

    toc_headings = []
    for line in lines:
        if line.startswith("## ") and not line.startswith("### "):
            toc_headings.append(line[3:].strip())
    add_toc(doc, toc_headings)

    in_code = False
    code_lines: list[str] = []
    i = 0
    active_number_id: int | None = None
    last_was_numbered = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("# "):
            last_was_numbered = False
            i += 1
            continue
        if stripped.startswith("```"):
            last_was_numbered = False
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped or stripped == "---":
            last_was_numbered = False
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            last_was_numbered = False
            rows, i = parse_table(lines, i)
            if rows:
                add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            last_was_numbered = False
            level = min(len(heading.group(1)) - 1, 3)
            title = heading.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title, base_size={1: 16, 2: 13, 3: 11.5}[level], color={1: BLUE_DARK, 2: BLUE, 3: CYAN}[level], force_bold=True)
            i += 1
            continue
        if stripped.startswith("> "):
            last_was_numbered = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(3)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHTER)
            p_pr.append(shd)
            add_inline(p, stripped[2:], base_size=9.3, color=MUTED)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            last_was_numbered = False
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if not last_was_numbered:
                active_number_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_number_id)
            add_inline(p, numbered.group(1))
            last_was_numbered = True
            i += 1
            continue
        p = doc.add_paragraph()
        last_was_numbered = False
        add_inline(p, stripped)
        i += 1

    props = doc.core_properties
    props.title = "DailyFlow 2.2 DeepSeek Harness 完整开发实施计划"
    props.subject = "AI Event Operator implementation handoff"
    props.author = "DailyFlow"
    props.keywords = "DailyFlow, DeepSeek Harness, AI Event Operator, Mindmap, Proposal"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
