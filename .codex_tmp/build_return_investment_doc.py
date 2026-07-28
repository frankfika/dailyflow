from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/fangchen/Baidu/GitHub/dailyflow/基金返投认定口径（讨论稿）.docx"

BODY_FONT = "Arial Unicode MS"
HEADING_FONT = "Arial Unicode MS"
ASCII_FONT = "Arial Unicode MS"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(100, 100, 100)


def set_run_font(run, cjk_font, size, *, bold=False, color=DARK):
    run.font.name = ASCII_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:cs"), ASCII_FONT)


def set_style_font(style, cjk_font, size, *, bold=False, color=DARK):
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:cs"), ASCII_FONT)


def set_keep_with_next(paragraph, value=True):
    ppr = paragraph._p.get_or_add_pPr()
    node = ppr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        ppr.append(node)
    elif not value and node is not None:
        ppr.remove(node)


def set_widow_control(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:widowControl")) is None:
        ppr.append(OxmlElement("w:widowControl"))


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])
    set_run_font(run, BODY_FONT, 9, color=MUTED)


def add_body(doc, text, *, first_line=True, keep_next=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(22) if first_line else Pt(0)
    p.paragraph_format.keep_together = False
    p.paragraph_format.keep_with_next = keep_next
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    set_widow_control(p)
    return p


def add_labeled_body(doc, label, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_run = p.add_run(label)
    set_run_font(label_run, HEADING_FONT, 11, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, BODY_FONT, 11)
    set_widow_control(p)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
set_style_font(normal, BODY_FONT, 11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

h1 = doc.styles["Heading 1"]
set_style_font(h1, HEADING_FONT, 16, bold=True, color=BLUE)
h1.paragraph_format.space_before = Pt(16)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.line_spacing = 1.0
h1.paragraph_format.keep_with_next = True
h1.paragraph_format.keep_together = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(20)
title.paragraph_format.line_spacing = 1.0
title_run = title.add_run("基金返投认定口径（讨论稿）")
set_run_font(title_run, HEADING_FONT, 22, bold=True, color=RGBColor(0, 0, 0))
set_keep_with_next(title)

add_body(
    doc,
    "结合北京开放传神科技有限公司（以下简称“开放传神”）作为本基金联合普通合伙人的产业资源及项目招引作用，建议返投按照以下口径认定：",
)

doc.add_heading("一、开放传神自身投入", level=1)
add_body(
    doc,
    "开放传神及其控股子公司在盐城实际发生的投入，按照实际投入金额的100%认定返投。",
)
add_body(
    doc,
    "认定范围包括但不限于实缴注册资本，研发及研发机构投入，实体项目、算力中心及相关基础设施建设投入，设备购置、场地租赁、本地用工等经营支出，向盐城本地企业采购产品或服务的实际支出，以及其他经各方认可的实际投入。",
)

doc.add_heading("二、开放传神招引项目", level=1)
add_body(
    doc,
    "由开放传神引进并落地盐城，但本基金没有投资的企业，按照其在盐城形成的实缴注册资本及其他实际投入的30%认定返投。",
)
add_body(
    doc,
    "其他实际投入包括但不限于研发及研发机构投入，厂房、生产基地及生产线建设投入，算力中心及相关基础设施建设投入，固定资产和设备购置投入，场地租赁、本地用工及本地采购等经营支出，其他在盐城形成的产业落地投入和本地经营支出，以及其他经各方认可的实际投入。",
)
add_body(doc, "上述实际投入不包括税款、政府收费及罚款等支出。")

doc.add_heading("三、基金投资项目", level=1)
add_body(
    doc,
    "本基金投资的企业在盐城设立子公司或实施产业项目的，按照其在盐城形成的实缴注册资本及其他实际投入，经去重后按照100%认定返投。",
)
add_body(
    doc,
    "其他实际投入包括但不限于厂房、生产基地及生产线建设投入，算力中心及相关基础设施建设投入，研发及研发机构投入，固定资产和设备购置投入，场地投入、本地用工及本地采购等经营支出，其他在盐城形成的产业落地投入和本地经营支出，以及其他经各方认可的实际投入。",
)
add_body(doc, "上述实际投入不包括税款、政府收费及罚款等支出。")
add_body(
    doc,
    "本基金投资的企业将总部或核心经营主体实质性迁入盐城的，按照本基金对该企业实际支付投资金额的1.5倍认定返投。",
)
add_body(
    doc,
    "同一笔资金不得重复计算。同一被投企业同时符合上述多种情形的，按照认定金额较高的一项计算；已经按照较低标准认定，后续达到较高标准的，只补充认定差额。",
)

doc.add_heading("四、统一认定原则", level=1)
add_labeled_body(
    doc,
    "（一）不重复计算。",
    "同一企业、同一笔资金或同一项投入不得重复计算。实缴注册资本后又使用该笔资金形成研发、建设、设备购置或其他支出的，只认定一次。",
)
add_labeled_body(
    doc,
    "（二）分类认定。",
    "开放传神及其控股子公司适用“开放传神自身投入”认定口径；由开放传神引进但本基金没有投资的企业，适用“开放传神招引项目”认定口径；本基金已经投资的企业，适用“基金投资项目”认定口径。",
)
add_labeled_body(
    doc,
    "（三）避免交叉。",
    "同一企业同时属于开放传神招引项目和基金投资项目的，按照基金投资项目认定，不再重复按照招引项目认定。",
)
add_labeled_body(
    doc,
    "（四）一次性足额认定。",
    "返投认定以相关投入实际发生并符合认定条件为准，与本基金分期出资、分期缴款安排无关。符合条件的返投应当一次性、足额认定，不得按照基金出资期数拆分认定。",
)
add_labeled_body(
    doc,
    "（五）累计结转使用。",
    "已经认定的返投金额统一计入累计返投金额，并可以结转用于满足基金后续各期及基金存续期内的返投要求。后续各期仅对累计返投金额进行核算，不对已经认定的返投重新拆分或重复认定。",
)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(0)
footer_p.paragraph_format.space_after = Pt(0)
add_page_number(footer_p)

doc.core_properties.title = "基金返投认定口径（讨论稿）"
doc.core_properties.subject = "基金返投认定方案"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""

doc.save(OUTPUT)
print(OUTPUT)
